from pypdf import PdfReader
from docx import Document

def extract_text(file):
    """
    Extract text from PDF or DOCX resume, handling tables in DOCX.
    """
    text = ""
    
    # Ensure pointer is at the start of the file
    file.seek(0)

    try:
        if file.name.lower().endswith(".pdf"):
            reader = PdfReader(file)
            pages = []

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)

            text = "\n".join(pages)

        elif file.name.lower().endswith(".docx"):
            doc = Document(file)
            full_text = []
            
            # 1. Extract from paragraphs (Standard text)
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
                
            # 2. Extract from tables (Commonly used for Resume Layouts)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Recursive text extraction from cells to capture paragraphs within cells
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                full_text.append(paragraph.text)
                            
            text = "\n".join(full_text)

    except Exception as e:
        print(f"Text extraction error: {str(e)}")
        # Return empty string will be caught by app.py as extraction failure
        return ""

    return text.strip()